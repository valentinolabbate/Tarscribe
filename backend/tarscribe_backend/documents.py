"""Text extraction for uploaded reference documents (PDF / DOCX / TXT / MD / HTML / EPUB).

Extracted plain text is fed into the same RAG pipeline as transcripts and
summaries (see ``rag.index_document``). Parser libraries are imported lazily so
a missing optional dependency degrades into a clear error instead of crashing
the whole backend at import time.
"""

from __future__ import annotations

from html.parser import HTMLParser
from pathlib import Path
from zipfile import BadZipFile, ZipFile

from .upload_security import DOCUMENT_UPLOAD_SUFFIXES

SUPPORTED_SUFFIXES = DOCUMENT_UPLOAD_SUFFIXES


class DocumentError(Exception):
    """Raised when a document cannot be read or contains no extractable text."""


def is_supported(filename: str) -> bool:
    return Path(filename).suffix.lower() in SUPPORTED_SUFFIXES


def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency missing
        raise DocumentError(
            "PDF-Unterstützung fehlt (pypdf nicht installiert)."
        ) from exc
    try:
        reader = PdfReader(str(path))
        return "\n\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception as exc:  # noqa: BLE001
        raise DocumentError(f"PDF konnte nicht gelesen werden: {exc}") from exc


def _extract_docx(path: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover - dependency missing
        raise DocumentError(
            "Word-Unterstützung fehlt (python-docx nicht installiert)."
        ) from exc
    try:
        document = docx.Document(str(path))
    except Exception as exc:  # noqa: BLE001
        raise DocumentError(f"Word-Dokument konnte nicht gelesen werden: {exc}") from exc
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" \t ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _extract_text_file(path: Path) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except (UnicodeDecodeError, UnicodeError):
            continue
    raise DocumentError("Textdatei-Encoding nicht erkannt.")


class _HTMLTextParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []

    def handle_data(self, data: str) -> None:
        text = data.strip()
        if text:
            self.parts.append(text)

    def text(self) -> str:
        return "\n".join(self.parts)


def _html_to_text(data: str) -> str:
    parser = _HTMLTextParser()
    parser.feed(data)
    return parser.text()


def _extract_html_file(path: Path) -> str:
    return _html_to_text(_extract_text_file(path))


def _extract_epub(path: Path) -> str:
    try:
        with ZipFile(path) as archive:
            parts: list[str] = []
            for name in archive.namelist():
                lower = name.lower()
                if lower.endswith((".html", ".htm", ".xhtml")):
                    raw = archive.read(name)
                    for encoding in ("utf-8", "utf-16", "latin-1"):
                        try:
                            parts.append(_html_to_text(raw.decode(encoding)))
                            break
                        except (UnicodeDecodeError, UnicodeError):
                            continue
            return "\n\n".join(part for part in parts if part.strip())
    except (BadZipFile, OSError) as exc:
        raise DocumentError(f"EPUB konnte nicht gelesen werden: {exc}") from exc


def extract_text(path: Path, content_type: str | None = None) -> str:
    """Return the plain-text content of a supported document file.

    Raises ``DocumentError`` for unsupported formats, read failures, or when no
    text can be extracted (e.g. a scanned PDF without an OCR layer).
    """
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text = _extract_pdf(path)
    elif suffix == ".docx":
        text = _extract_docx(path)
    elif suffix in {".txt", ".md", ".markdown", ".text"}:
        text = _extract_text_file(path)
    elif suffix in {".html", ".htm"}:
        text = _extract_html_file(path)
    elif suffix == ".epub":
        text = _extract_epub(path)
    else:
        raise DocumentError(f"Nicht unterstütztes Dateiformat: {suffix or '(keine Endung)'}")

    text = (text or "").strip()
    if not text:
        raise DocumentError(
            "Keine Textinhalte gefunden. Bei gescannten PDFs fehlt eine Textebene (OCR)."
        )
    return text
